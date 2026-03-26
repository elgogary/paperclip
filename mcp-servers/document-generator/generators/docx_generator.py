"""DOCX generator — populates a branded .docx template with structured content."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _apply_default_style(doc: Document, brand: dict):
    style = doc.styles["Normal"]
    style.font.name = brand.get("font", "Calibri")
    style.font.size = Pt(11)
    style.font.color.rgb = _hex_to_rgb(brand.get("text", "#333333"))
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = brand.get("font", "Calibri")
        h.font.color.rgb = _hex_to_rgb(brand.get("primary", "#1B3A57"))


def _add_title_page(doc: Document, data: dict, brand: dict):
    doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.get("title", "Document"))
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = _hex_to_rgb(brand.get("primary", "#1B3A57"))

    if data.get("subtitle"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(data["subtitle"])
        run.font.size = Pt(16)
        run.font.color.rgb = _hex_to_rgb(brand.get("accent", "#4A6F8F"))

    meta = data.get("metadata", {})
    if meta:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for key, val in meta.items():
            p.add_run(f"{key}: {val}\n").font.size = Pt(11)

    doc.add_page_break()


def _add_section(doc: Document, section: dict, brand: dict):
    level = section.get("level", 1)
    if section.get("heading"):
        doc.add_heading(section["heading"], level=min(level, 3))

    for para in section.get("paragraphs", []):
        doc.add_paragraph(para)

    for bullet in section.get("bullets", []):
        doc.add_paragraph(f"  {bullet}", style="List Bullet")

    tbl_data = section.get("table")
    if tbl_data:
        headers = tbl_data.get("headers", [])
        rows = tbl_data.get("rows", [])
        if headers:
            table = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
            for row_data in rows:
                row = table.add_row()
                for i, val in enumerate(row_data):
                    if i < len(row.cells):
                        row.cells[i].text = str(val)

    img = section.get("image")
    if img and Path(img).exists():
        doc.add_picture(img, width=Inches(5.5))

    if section.get("page_break"):
        doc.add_page_break()


def generate_docx(template_path: Path | None, data: dict, output_path: Path) -> dict:
    """Generate a .docx from template + content data.

    Returns dict with metadata about the generated document.
    """
    brand = data.get("brand", {})

    if template_path and template_path.exists():
        doc = Document(str(template_path))
    else:
        doc = Document()
        _apply_default_style(doc, brand)

    if data.get("title"):
        _add_title_page(doc, data, brand)

    sections = data.get("sections", [])
    for section in sections:
        _add_section(doc, section, brand)

    doc.save(str(output_path))

    return {
        "sections": len(sections),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }
